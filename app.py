import os
import re
import io
import time
import tempfile
import zipfile as _zip
from datetime import datetime
from pathlib import Path
from typing import Tuple, Dict, Any, List

import requests
import streamlit as st
import pandas as pd
from bs4 import BeautifulSoup
import PyPDF2
import docx
import ebooklib
from ebooklib import epub

# === Google GenAI (new unified SDK) ===
from google import genai
from google.genai import types
from google.genai import errors as genai_errors

# =========================
# Page & Globals
# =========================
st.set_page_config(
    page_title="Samawy Book Blurb Writer — Streamlit",
    page_icon="📖",
    layout="wide",
)

SUPPORTED_MODELS = ["gemini-2.0-flash", "gemini-2.0-flash-lite", "gemini-2.5-pro"]
PRICING_URL = "https://ai.google.dev/gemini-api/docs/pricing"  # official pricing page

BOOK_CATEGORIES = [
    "التقنية والكمبيوتر", "القواميس والموسوعات", "معلومات عامة", "العلوم الاجتماعية والسياسية",
    "التراجم والسير", "التاريخ والجغرافيا", "الإدارة والأعمال", "القصة والرواية", "القانون",
    "العلوم والرياضيات", "الهوايات والأشغال اليدوية", "تعليم اللغات", "هندسة العمارة والتصميم",
    "الطبخ", "المجلات", "السفر والخرائط", "الفلسفة والفكر", "المقررات والمناهج", "كتب الأطفال",
    "المرأة والأسرة", "الصحة العامة والتغذية والحمية", "الكتب المدرسية", "الكتب الطبية",
    "الأدب والشعر", "الطبيعة والزراعة وعلم الحيوان", "تطوير الذات", "العناية بالطفل",
    "التربية والتعليم", "كتب الهندسة", "الكتب الإسلامية والدينية"
]

# Session vars
if "last_stats" not in st.session_state:
    st.session_state.last_stats = {}  # holds latest run stats
if "op_log" not in st.session_state:
    st.session_state.op_log: List[Dict[str, Any]] = []  # append a dict per operation
if "live_prices" not in st.session_state:
    st.session_state.live_prices = {}  # filled at app start
if "price_in" not in st.session_state:
    st.session_state.price_in = 0.0
if "price_out" not in st.session_state:
    st.session_state.price_out = 0.0

# =========================
# Helpers
# =========================
def get_client():
    """Initialize the Gemini client using Streamlit secrets or env var."""
    api_key = (
        st.secrets.get("GEMINI_API_KEY")
        or os.environ.get("GEMINI_API_KEY")
        or os.environ.get("GOOGLE_API_KEY")
    )
    if not api_key:
        st.error("No API key found. Add `GEMINI_API_KEY` in App → Settings → Secrets (or set env var).")
        st.stop()
    try:
        return genai.Client(api_key=api_key)
    except Exception as e:
        st.error(f"Failed to initialize Gemini client: {e}")
        st.stop()

@st.cache_data(show_spinner=False)
def list_available_models_cached() -> List[str]:
    """Return model IDs available to this key. Cached to reduce latency."""
    try:
        client = get_client()
        models = client.models.list()
        return sorted({m.name.split("/")[-1] for m in models})
    except Exception:
        return []

def resolve_model_id(preferred: str) -> str:
    """Return a working model id or a safe default from SUPPORTED_MODELS."""
    available = set(list_available_models_cached())
    if preferred in available:
        return preferred
    for m in SUPPORTED_MODELS:
        if m in available:
            return m
    return preferred  # best-effort

def _usage_counts(resp) -> Tuple[int, int]:
    """Safely extract token counts from response.usage_metadata."""
    usage = getattr(resp, "usage_metadata", None)
    if not usage:
        return 0, 0
    try:
        pt = int(getattr(usage, "input_tokens", 0) or 0)
        ot = int(getattr(usage, "output_tokens", 0) or 0)
        return pt, ot
    except Exception:
        return 0, 0

def _model_name_from_response(resp: object, fallback: str) -> str:
    """Return a clean model id from response fields, or fallback."""
    raw = getattr(resp, "model_version", None) or getattr(resp, "model", None) or fallback
    try:
        return str(raw).split("/")[-1]
    except Exception:
        return fallback

def estimate_tokens_from_chars(s: str) -> int:
    """Heuristic: ~1 token per ~4 characters (Google guidance)."""
    if not s:
        return 0
    return max(0, int(len(s.strip()) / 4))

def calculate_cost(prompt_tokens: int, output_tokens: int, price_in: float, price_out: float) -> float:
    return ((prompt_tokens/1_000_000) * price_in) + ((output_tokens/1_000_000) * price_out)

# =========================
# Live Pricing Fetch & Parse
# =========================
@st.cache_data(show_spinner=False, ttl=3600)
def fetch_pricing_html(url: str) -> str:
    """Fetch the pricing page HTML (cached for 1 hour)."""
    resp = requests.get(url, timeout=20)
    resp.raise_for_status()
    return resp.text

def _money_to_float(s: str) -> float:
    m = re.search(r"\$?\s*([0-9]+(?:\.[0-9]+)?)", s)
    return float(m.group(1)) if m else 0.0

def parse_gemini_pricing(html: str) -> Dict[str, Dict[str, float]]:
    """
    Parse the official pricing page to extract per-1M token input/output prices
    for key models we care about. The page can change; we use robust, text-first
    heuristics and prefer 'text/image/video' input and 'output (including thinking)'.
    Returns: { model_id: {"input": float, "output": float} }
    """
    soup = BeautifulSoup(html, "lxml")
    text = soup.get_text("\n", strip=True)
    lo = text.lower()

    # Candidate labels to anchor sections
    anchors = {
        "gemini-2.0-flash": [
            "gemini 2.0 flash", "2.0 flash"
        ],
        "gemini-2.0-flash-lite": [
            "gemini 2.0 flash lite", "2.0 flash lite", "flash-lite"
        ],
        "gemini-2.5-pro": [
            "gemini 2.5 pro", "2.5 pro"
        ],
        # You can extend with "gemini 2.5 flash" or "2.5 flash live" if you plan to use them.
    }

    # For each model, locate the nearest "input price" and "output price" lines after the anchor.
    prices: Dict[str, Dict[str, float]] = {}
    for model, keys in anchors.items():
        idx = -1
        for k in keys:
            idx = lo.find(k)
            if idx != -1:
                break
        if idx == -1:
            continue  # not found on page text

        # Search a local window after the anchor
        window = lo[idx: idx + 2500]  # look ahead 2500 chars
        # Extract lines for clarity
        lines = window.splitlines()

        input_price = 0.0
        output_price = 0.0

        # Heuristics: find "input price" line and "output price" line
        for i, line in enumerate(lines):
            if "input price" in line:
                # Prefer the first $ amount on same line or the next line(s)
                seg = lines[i]
                # If input covers multiple modalities on separate lines, we prefer text/image/video one
                # Try this line first:
                val = _money_to_float(seg)
                if val == 0.0 and i+1 < len(lines):
                    # check a couple of following lines
                    for j in range(1, 5):
                        if i + j < len(lines):
                            val = _money_to_float(lines[i + j])
                            if val:
                                break
                if val:
                    input_price = val

            if "output price" in line:
                seg = lines[i]
                val = _money_to_float(seg)
                if val == 0.0 and i+1 < len(lines):
                    for j in range(1, 5):
                        if i + j < len(lines):
                            val = _money_to_float(lines[i + j])
                            if val:
                                break
                if val:
                    output_price = val

        # If still zero (page structure is different), try broader regex around the window
        if input_price == 0.0:
            m = re.search(r"input price.*?\$?\s*([0-9]+(?:\.[0-9]+)?)", window, flags=re.S)
            if m:
                input_price = float(m.group(1))
        if output_price == 0.0:
            m = re.search(r"output price.*?\$?\s*([0-9]+(?:\.[0-9]+)?)", window, flags=re.S)
            if m:
                output_price = float(m.group(1))

        if input_price or output_price:
            prices[model] = {"input": input_price, "output": output_price}

    return prices

def get_live_prices() -> Dict[str, Dict[str, float]]:
    """Fetch and parse live pricing; returns dict. Falls back to previous cached values."""
    try:
        html = fetch_pricing_html(PRICING_URL)
        parsed = parse_gemini_pricing(html)
        # Keep prior values if some models missing
        merged = {**st.session_state.live_prices, **parsed} if st.session_state.live_prices else parsed
        return merged or {}
    except Exception as e:
        st.warning(f"Could not fetch live pricing: {e}")
        return st.session_state.live_prices or {}

def pick_default_prices(prices: Dict[str, Dict[str, float]], model: str) -> Tuple[float, float]:
    """
    Choose default (input, output) from live prices for a given model id.
    If a model isn't found, return (0.0, 0.0).
    """
    info = prices.get(model) or {}
    return float(info.get("input", 0.0) or 0.0), float(info.get("output", 0.0) or 0.0)

# =========================
# Extraction
# =========================
def extract_text_from_pdf(file_bytes: bytes) -> str:
    text = ""
    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    for page in reader.pages:
        t = page.extract_text()
        if t:
            text += t + "\n"
    return text.strip()

def extract_text_from_docx(file_bytes: bytes) -> str:
    with io.BytesIO(file_bytes) as buff:
        doc = docx.Document(buff)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts).strip()

def extract_text_from_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore").strip()

def extract_text_from_epub(file_bytes: bytes) -> str:
    # ebooklib.read_epub needs a real path — write to temp file.
    with tempfile.NamedTemporaryFile(suffix=".epub", delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name
    try:
        book = epub.read_epub(tmp_path)
        text = ""
        for it in book.get_items():
            if it.get_type() == ebooklib.ITEM_DOCUMENT:
                soup = BeautifulSoup(it.get_content(), "html.parser")
                text += soup.get_text(separator="\n") + "\n"
        return text.strip()
    finally:
        try:
            os.remove(tmp_path)
        except Exception:
            pass

def extract_text_from_indd(file_bytes: bytes) -> str:
    # Heuristic: scrape long printable runs
    matches = re.findall(rb'[\x20-\x7E\xA0-\xFF]{10,}', file_bytes)
    return ' '.join([m.decode('utf-8', errors='ignore') for m in matches]).strip()

def extract_text(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    data = uploaded_file.read()
    if name.endswith(".pdf"):
        return extract_text_from_pdf(data)
    if name.endswith(".docx") or name.endswith(".doc"):
        return extract_text_from_docx(data)
    if name.endswith(".txt"):
        return extract_text_from_txt(data)
    if name.endswith(".epub"):
        return extract_text_from_epub(data)
    if name.endswith(".indd"):
        return extract_text_from_indd(data)
    raise ValueError(f"Unsupported format: {name}")

def clean_text(text: str) -> str:
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^\w\s\.,!?;:\'"()\-\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]', '', text)
    return re.sub(r'\s+', ' ', text).strip()

def chunk_text(text: str, max_words: int = 1500, start_percentage: float = 0.4) -> Tuple[str, str]:
    words = text.split()
    n = len(words)
    if n <= max_words:
        return text, f"The complete available text ({n} words)."
    num_start = int(max_words * start_percentage)
    num_rem = max_words - num_start
    start_words = words[:num_start]
    remaining = words[num_start:]
    if num_rem <= 0 or not remaining:
        final_words = start_words
        descr = f"A sample consisting of the first {len(start_words)} words. Total sample: {len(final_words)} words."
    else:
        if len(remaining) <= num_rem:
            final_words = start_words + remaining
            descr = f"A sample including the first {len(start_words)} words and all remaining {len(remaining)} words. Total: {len(final_words)}."
        else:
            interval = max(1, len(remaining) // num_rem)
            sampled = [remaining[i] for i in range(0, len(remaining), interval)][:num_rem]
            final_words = start_words + sampled
            descr = f"A sample including the first {len(start_words)} words and {len(sampled)} words sampled from the rest. Total: {len(final_words)}."
    return ' '.join(final_words).strip(), descr

# =========================
# Prompt Builders
# =========================
def build_blurb_prompt(text_chunk: str, chunk_description: str) -> str:
    return f"""
بالاستناد إلى المقطع التالي المقتبس من كتاب، أَنتِج نبذة تعريفية قصيرة بالعربية (١٦٠–٥٠٠ حرف).

- نبرة أدبية قريبة من أسلوب النص.
- تعكس جوهر الموضوع أو النبرة الشعورية.
- عربية سليمة دون دعاية صريحة أو أوامر (مثل: "اقرأ").
- دون حشو أو تكرار. علامات تشكيل فقط عند الضرورة.
- فضّل الجمل الفعلية. لا تبدأ بشبه جملة.

معلومة عن المقطع: {chunk_description}

النص (حتى ٣٠٠٠ حرف):
{text_chunk[:3000]}

أخرج "النبذة" فقط.
""".strip()

def build_category_prompt(text_chunk: str, chunk_description: str, categories: List[str]) -> str:
    cats_str = ", ".join(categories)
    return f"""
Based on the following Arabic book text, pick ONE category from this list:
{cats_str}

Info about the sample: "{chunk_description}"

Text (<=3000 chars):
{text_chunk[:3000]}

Return only the category name, nothing else.
""".strip()

# =========================
# Cached model calls
# =========================
def _gen_content(model: str, prompt: str):
    client = get_client()
    try:
        return client.models.generate_content(
            model=model,
            contents=[prompt],
            config=types.GenerateContentConfig(),
        )
    except genai_errors.APIError as e:
        if e.code == 404:
            fallback = resolve_model_id(model)
            if fallback != model:
                return client.models.generate_content(
                    model=fallback, contents=[prompt], config=types.GenerateContentConfig()
                )
        raise

@st.cache_data(show_spinner=False)
def generate_blurb_cached(model: str, text_chunk: str, chunk_description: str) -> Dict[str, Any]:
    """Cached blurb call. Returns dict with text, tokens, used model, prompt."""
    prompt = build_blurb_prompt(text_chunk, chunk_description)
    resp = _gen_content(model, prompt)
    blurb_text = (getattr(resp, "text", None) or "").strip()
    pt, ot = _usage_counts(resp)
    used_model = _model_name_from_response(resp, model)
    if not blurb_text:
        blurb_text = "Error generating blurb."
    return {
        "blurb": blurb_text.ljust(160),
        "prompt_tokens": pt,
        "output_tokens": ot,
        "used_model": used_model,
        "prompt": prompt,
    }

@st.cache_data(show_spinner=False)
def categorize_book_cached(model: str, text_chunk: str, chunk_description: str) -> Dict[str, Any]:
    """Cached categorization call. Returns dict with category, tokens, used model, prompt."""
    prompt = build_category_prompt(text_chunk, chunk_description, BOOK_CATEGORIES)
    resp = _gen_content(model, prompt)
    cat_text = (getattr(resp, "text", None) or "").strip()
    pt, ot = _usage_counts(resp)
    used_model = _model_name_from_response(resp, model)
    if cat_text not in BOOK_CATEGORIES:
        # fuzzy fallback
        match = next((c for c in BOOK_CATEGORIES if cat_text.lower() in c.lower() or c.lower() in cat_text.lower()), "القصة والرواية")
        cat_text = match
    return {
        "category": cat_text,
        "prompt_tokens": pt,
        "output_tokens": ot,
        "used_model": used_model,
        "prompt": prompt,
    }

# =========================
# Sidebar (with live pricing)
# =========================
st.title("📖 Samawy Book Blurb Writer — Streamlit Edition")
st.caption("AI-Powered Arabic blurbs & categorization (Gemini 2.x/2.5).")

with st.sidebar:
    st.header("🔑 AI Configuration")
    chosen_model = st.selectbox(
        "Model",
        options=SUPPORTED_MODELS,
        index=0,
        help="Use 2.x/2.5 models to avoid legacy 404s.",
    )
    if st.button("List Available Models"):
        st.write(list_available_models_cached())

    st.divider()
    st.subheader("💸 Live Pricing (per 1M tokens)")

    # Fetch live prices on first load or refresh
    col_lp1, col_lp2 = st.columns([1, 1])
    with col_lp1:
        if st.button("Refresh from Google"):
            st.session_state.live_prices = get_live_prices()
    # If empty, try to fill at first run
    if not st.session_state.live_prices:
        st.session_state.live_prices = get_live_prices()

    # Show a small table of fetched prices
    if st.session_state.live_prices:
        df_prices = pd.DataFrame([
            {"Model": k, "Input": v.get("input", 0.0), "Output": v.get("output", 0.0)}
            for k, v in st.session_state.live_prices.items()
        ])
        st.dataframe(df_prices, use_container_width=True, height=180)
    else:
        st.info("No live prices detected yet. You can still set prices manually below.")

    # Prefill the editable inputs from live price of chosen model (unless user changed them already)
    live_in, live_out = pick_default_prices(st.session_state.live_prices, chosen_model)

    def _prefill_once():
        if st.session_state.price_in == 0.0 and live_in:
            st.session_state.price_in = float(live_in)
        if st.session_state.price_out == 0.0 and live_out:
            st.session_state.price_out = float(live_out)

    _prefill_once()

    st.number_input(
        "Input price per 1M tokens (USD)",
        min_value=0.0,
        value=float(st.session_state.price_in),
        step=0.05,
        key="price_in"
    )
    st.number_input(
        "Output price per 1M tokens (USD)",
        min_value=0.0,
        value=float(st.session_state.price_out),
        step=0.05,
        key="price_out"
    )

    if st.button("Reset to live price for selected model"):
        st.session_state.price_in, st.session_state.price_out = pick_default_prices(st.session_state.live_prices, chosen_model)
        st.experimental_rerun()

# =========================
# Tabs: Single | Bulk | Stats | Log
# =========================
tab_single, tab_bulk, tab_stats, tab_log = st.tabs(["Single File", "Bulk", "Stats", "Log"])

# ---------- Single File ----------
with tab_single:
    st.subheader("Single File Analysis")
    up = st.file_uploader(
        "Upload one book file (.pdf, .docx/.doc, .txt, .epub, .indd)",
        type=["pdf", "docx", "doc", "txt", "epub", "indd"],
    )
    if up is not None:
        with st.spinner("Extracting text..."):
            try:
                raw = extract_text(up)
            except Exception as e:
                st.error(f"Extraction failed: {e}")
                st.stop()
        if not raw:
            st.warning("No text found in this file.")
            st.stop()

        clean = clean_text(raw)
        chunk, descr = chunk_text(clean, max_words=1500, start_percentage=0.4)

        col_meta, col_words = st.columns(2)
        with col_meta:
            st.markdown("**AI Input Description**")
            st.info(descr)
        with col_words:
            st.markdown("**AI Input Word Count**")
            st.metric(label="Words sent to AI", value=len(chunk.split()))

        if st.button("Generate Blurb & Category"):
            try:
                with st.spinner("Generating blurb..."):
                    blurb_res = generate_blurb_cached(chosen_model, chunk, descr)
                with st.spinner("Categorizing..."):
                    cat_res = categorize_book_cached(chosen_model, chunk, descr)
            except genai_errors.APIError as e:
                st.error(f"Gemini API error [{e.code}]: {e.message}")
                st.stop()
            except Exception as e:
                st.error(f"Error: {e}")
                st.stop()

            # Show results in this tab
            st.write("### 📝 Blurb")
            st.text_area("Generated Blurb", value=blurb_res["blurb"], height=200)
            st.caption(f"Characters: {len(blurb_res['blurb'])}")

            st.write("### 📚 Category")
            st.success(cat_res["category"])

            # Build robust stats (real or estimated)
            real_prompt = int(blurb_res["prompt_tokens"] or 0) + int(cat_res["prompt_tokens"] or 0)
            real_output = int(blurb_res["output_tokens"] or 0) + int(cat_res["output_tokens"] or 0)

            if real_prompt == 0 or real_output == 0:
                est_p_blurb = estimate_tokens_from_chars(blurb_res["prompt"])
                est_p_cat   = estimate_tokens_from_chars(cat_res["prompt"])
                est_c_blurb = estimate_tokens_from_chars(blurb_res["blurb"])
                est_c_cat   = estimate_tokens_from_chars(cat_res["category"])
                prompt_tokens = real_prompt if real_prompt > 0 else (est_p_blurb + est_p_cat)
                output_tokens = real_output if real_output > 0 else (est_c_blurb + est_c_cat)
                show_p_blurb = blurb_res["prompt_tokens"] if blurb_res["prompt_tokens"] else f"~{est_p_blurb}"
                show_c_blurb = blurb_res["output_tokens"] if blurb_res["output_tokens"] else f"~{est_c_blurb}"
                show_p_cat   = cat_res["prompt_tokens"]   if cat_res["prompt_tokens"]   else f"~{est_p_cat}"
                show_c_cat   = cat_res["output_tokens"]   if cat_res["output_tokens"]   else f"~{est_c_cat}"
            else:
                prompt_tokens = real_prompt
                output_tokens = real_output
                show_p_blurb  = blurb_res["prompt_tokens"]
                show_c_blurb  = blurb_res["output_tokens"]
                show_p_cat    = cat_res["prompt_tokens"]
                show_c_cat    = cat_res["output_tokens"]

            cleaned_words = len(clean.split())
            chunk_words   = len(chunk.split())
            est_cost = calculate_cost(prompt_tokens, output_tokens, float(st.session_state.price_in), float(st.session_state.price_out))

            # Update global "last_stats"
            st.session_state.last_stats = {
                "timestamp": datetime.now().isoformat(timespec="seconds"),
                "file": up.name,
                "model_blurb": blurb_res["used_model"],
                "model_cat": cat_res["used_model"],
                "cleaned_words": cleaned_words,
                "chunk_words": chunk_words,
                "p_blurb": show_p_blurb,
                "c_blurb": show_c_blurb,
                "p_cat": show_p_cat,
                "c_cat": show_c_cat,
                "total_tokens": prompt_tokens + output_tokens,
                "prompt_tokens": prompt_tokens,
                "output_tokens": output_tokens,
                "estimated_cost": est_cost,
            }

            # Append to log
            st.session_state.op_log.append({
                "timestamp": st.session_state.last_stats["timestamp"],
                "mode": "single",
                "file": up.name,
                "model_blurb": blurb_res["used_model"],
                "model_cat": cat_res["used_model"],
                "blurb": blurb_res["blurb"],
                "category": cat_res["category"],
                "ai_input_words": chunk_words,
                "p_blurb": show_p_blurb, "c_blurb": show_c_blurb,
                "p_cat": show_p_cat, "c_cat": show_c_cat,
                "total_tokens": st.session_state.last_stats["total_tokens"],
                "estimated_cost": est_cost,
            })

            st.success("Done! Check the **Stats** tab for detailed metrics and the **Log** tab for history.")

# ---------- Bulk ----------
with tab_bulk:
    st.subheader("Bulk Analysis (Upload many files)")
    save_samples = st.checkbox("Save the AI input text samples to a downloadable .zip")
    ups = st.file_uploader(
        "Upload multiple files", type=["pdf","docx","doc","txt","epub","indd"], accept_multiple_files=True
    )
    if ups:
        rows = []
        sample_files = []
        for f in ups:
            name = f.name
            try:
                raw = extract_text(f)
                if not raw:
                    rows.append({"File": name, "Status": "Error: No text", "Blurb": "", "Category": ""})
                    continue
                clean_b = clean_text(raw)
                chunk_b, descr_b = chunk_text(clean_b)

                # Cached calls
                blurb_res = generate_blurb_cached(chosen_model, chunk_b, descr_b)
                cat_res   = categorize_book_cached(chosen_model, chunk_b, descr_b)

                # Real or estimated tokens
                real_p = int(blurb_res["prompt_tokens"] or 0) + int(cat_res["prompt_tokens"] or 0)
                real_o = int(blurb_res["output_tokens"] or 0) + int(cat_res["output_tokens"] or 0)

                if real_p == 0 or real_o == 0:
                    est_p1 = estimate_tokens_from_chars(blurb_res["prompt"])
                    est_p2 = estimate_tokens_from_chars(cat_res["prompt"])
                    est_c1 = estimate_tokens_from_chars(blurb_res["blurb"])
                    est_c2 = estimate_tokens_from_chars(cat_res["category"])
                    p_total = real_p if real_p > 0 else (est_p1 + est_p2)
                    o_total = real_o if real_o > 0 else (est_c1 + est_c2)
                    disp_p1 = blurb_res["prompt_tokens"] if blurb_res["prompt_tokens"] else f"~{est_p1}"
                    disp_c1 = blurb_res["output_tokens"] if blurb_res["output_tokens"] else f"~{est_c1}"
                    disp_p2 = cat_res["prompt_tokens"]   if cat_res["prompt_tokens"]   else f"~{est_p2}"
                    disp_c2 = cat_res["output_tokens"]   if cat_res["output_tokens"]   else f"~{est_c2}"
                else:
                    p_total = real_p
                    o_total = real_o
                    disp_p1 = blurb_res["prompt_tokens"]
                    disp_c1 = blurb_res["output_tokens"]
                    disp_p2 = cat_res["prompt_tokens"]
                    disp_c2 = cat_res["output_tokens"]

                est_cost = calculate_cost(p_total, o_total, float(st.session_state.price_in), float(st.session_state.price_out))
                ai_input_words = len(chunk_b.split())

                rows.append({
                    "File": name,
                    "Status": "Success",
                    "Blurb": blurb_res["blurb"],
                    "Category": cat_res["category"],
                    "AI Input Words": ai_input_words,
                    "Blurb Prompt": disp_p1, "Blurb Output": disp_c1,
                    "Cat Prompt": disp_p2,   "Cat Output":  disp_c2,
                    "Model (Blurb)": blurb_res["used_model"],
                    "Model (Cat)": cat_res["used_model"],
                    "Total Tokens": p_total + o_total,
                    "Estimated Cost": f"${est_cost:.6f}",
                })

                # Save samples for download
                if save_samples:
                    sample_files.append((Path(name).stem + "_sample.txt", chunk_b))

                # Log each bulk item
                st.session_state.op_log.append({
                    "timestamp": datetime.now().isoformat(timespec="seconds"),
                    "mode": "bulk",
                    "file": name,
                    "model_blurb": blurb_res["used_model"],
                    "model_cat": cat_res["used_model"],
                    "blurb": blurb_res["blurb"],
                    "category": cat_res["category"],
                    "ai_input_words": ai_input_words,
                    "p_blurb": disp_p1, "c_blurb": disp_c1,
                    "p_cat": disp_p2, "c_cat": disp_c2,
                    "total_tokens": p_total + o_total,
                    "estimated_cost": est_cost,
                })

            except Exception as e:
                rows.append({"File": name, "Status": f"Error: {e}", "Blurb": "", "Category": ""})

        df = pd.DataFrame(rows)
        st.dataframe(df, use_container_width=True)

        # Download Excel
        if not df.empty:
            out = io.BytesIO()
            with pd.ExcelWriter(out, engine="openpyxl") as w:
                df.to_excel(w, index=False, sheet_name="Results")
            st.download_button("⬇️ Download Excel", data=out.getvalue(), file_name="samawy_blurb_results.xlsx")

        # Download text samples
        if save_samples and sample_files:
            zbuf = io.BytesIO()
            with _zip.ZipFile(zbuf, "w", _zip.ZIP_DEFLATED) as zf:
                for fname, content in sample_files:
                    zf.writestr(fname, content)
            st.download_button("⬇️ Download Text Samples (.zip)", data=zbuf.getvalue(), file_name="ai_input_samples.zip")

        st.info("Bulk complete. See **Stats** for the latest single-run metrics (if any) and **Log** for full history.")

# ---------- Stats (own tab) ----------
with tab_stats:
    st.subheader("📈 Latest Stats")
    ls = st.session_state.last_stats
    if not ls:
        st.info("Run a Single or Bulk operation to populate Stats.")
    else:
        stats_lines = [
            f"Timestamp: {ls['timestamp']}",
            f"File: {ls['file']}",
            f"Model (blurb): {ls['model_blurb']}",
            f"Model (category): {ls['model_cat']}",
            f"Total Cleaned Words: {ls['cleaned_words']:,}",
            f"AI Input Words: {ls['chunk_words']:,}",
            f"Blurb (Prompt/Output): {ls['p_blurb']} / {ls['c_blurb']} tokens",
            f"Category (Prompt/Output): {ls['p_cat']} / {ls['c_cat']} tokens",
            f"Total Tokens Used: {ls['total_tokens']:,}",
            f"Estimated Cost: ${ls['estimated_cost']:.6f}",
        ]
        st.code("\n".join(stats_lines))

# ---------- Log (cached operations) ----------
with tab_log:
    st.subheader("🧾 Operation Log (Cached)")
    if not st.session_state.op_log:
        st.info("No operations yet. Run Single or Bulk to populate the log.")
    else:
        df_log = pd.DataFrame(st.session_state.op_log)
        st.dataframe(df_log, use_container_width=True)
        out = io.BytesIO()
        with pd.ExcelWriter(out, engine="openpyxl") as w:
            df_log.to_excel(w, index=False, sheet_name="Log")
        st.download_button("⬇️ Download Log (Excel)", data=out.getvalue(), file_name="samawy_blurb_log.xlsx")
